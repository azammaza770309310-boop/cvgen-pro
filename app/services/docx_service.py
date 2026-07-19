"""DOCX export service — python-docx.

Generates a valid DOCX from the SAME ResumeData. Supports EN/AR/bilingual with
RTL paragraph direction for Arabic.
"""
from __future__ import annotations

import io
import logging

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

from app.models.resume import ResumeData
from app.services.resume_normalizer import normalize_resume_data
from app.utils.arabic import contains_arabic

logger = logging.getLogger("cvgen.docx")


def _set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def _add_heading(doc, text, level=1, rtl=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(22)
    elif level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    else:
        run.font.size = Pt(12)
    if rtl:
        _set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return p


def _add_para(doc, text, rtl=False, italic=False):
    if not text:
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    if italic:
        run.italic = True
    if rtl:
        _set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return p


def _add_bullet(doc, text, rtl=False):
    if not text:
        return
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    if rtl:
        _set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _lang_fields(resume: ResumeData, lang: str):
    is_ar = lang == "ar"
    name = resume.personal.name_ar if is_ar else (resume.personal.name_en or resume.personal.name)
    title = resume.personal.title_ar if is_ar else (resume.personal.title_en or resume.personal.title)
    summary = resume.summary_text(lang)
    return name, title, summary, is_ar


def _write_section(doc, resume: ResumeData, lang: str):
    name, title, summary, is_ar = _lang_fields(resume, lang)
    if name:
        _add_heading(doc, name, level=0, rtl=is_ar)
    if title:
        _add_para(doc, title, rtl=is_ar, italic=True)
    contact_parts = []
    for v in (resume.personal.email, resume.personal.phone, resume.personal.location, resume.personal.linkedin, resume.personal.website):
        if v:
            contact_parts.append(v)
    if contact_parts:
        _add_para(doc, " · ".join(contact_parts), rtl=is_ar)
    doc.add_paragraph()

    if summary:
        _add_heading(doc, "ملخص" if is_ar else "Summary", level=1, rtl=is_ar)
        _add_para(doc, summary, rtl=is_ar)

    if resume.experience:
        _add_heading(doc, "الخبرات" if is_ar else "Experience", level=1, rtl=is_ar)
        for e in resume.experience:
            t = e.title_ar if is_ar else (e.title_en or e.title)
            c = e.company_ar if is_ar else (e.company_en or e.company)
            d = ""
            if e.start_date or e.end_date:
                d = f"{e.start_date} – {e.end_date or ('Present' if e.current else '')}".strip(" –")
            if t:
                _add_heading(doc, t, level=2, rtl=is_ar)
            if c or d:
                _add_para(doc, f"{c}  {('· '+d) if d else ''}".strip(), rtl=is_ar, italic=True)
            bullets = e.bullets_ar if is_ar else (e.bullets_en or e.bullets)
            for b in bullets:
                _add_bullet(doc, b, rtl=is_ar)

    if resume.education:
        _add_heading(doc, "التعليم" if is_ar else "Education", level=1, rtl=is_ar)
        for ed in resume.education:
            deg = ed.degree_ar if is_ar else (ed.degree_en or ed.degree)
            inst = ed.institution_ar if is_ar else (ed.institution_en or ed.institution)
            if deg:
                _add_heading(doc, deg, level=2, rtl=is_ar)
            if inst or ed.year:
                _add_para(doc, f"{inst}  {('· '+ed.year) if ed.year else ''}".strip(), rtl=is_ar, italic=True)

    all_skills = resume.skills + resume.technical_skills + resume.soft_skills
    if all_skills:
        _add_heading(doc, "المهارات" if is_ar else "Skills", level=1, rtl=is_ar)
        _add_para(doc, ", ".join(all_skills), rtl=is_ar)

    if resume.courses:
        _add_heading(doc, "الدورات" if is_ar else "Courses", level=1, rtl=is_ar)
        _add_para(doc, ", ".join(resume.courses), rtl=is_ar)

    if resume.certifications:
        _add_heading(doc, "الشهادات" if is_ar else "Certifications", level=1, rtl=is_ar)
        _add_para(doc, ", ".join(c.name for c in resume.certifications), rtl=is_ar)

    if resume.languages:
        _add_heading(doc, "اللغات" if is_ar else "Languages", level=1, rtl=is_ar)
        _add_para(doc, ", ".join(f"{l.name} ({l.level})" if l.level else l.name for l in resume.languages), rtl=is_ar)


def export_docx(resume_data: ResumeData | dict) -> bytes:
    if isinstance(resume_data, dict):
        resume = normalize_resume_data(resume_data)
    else:
        resume = resume_data
    doc = Document()
    # default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lang = resume.lang or "en"
    if lang == "bilingual":
        _write_section(doc, resume, "en")
        doc.add_page_break()
        _write_section(doc, resume, "ar")
    else:
        _write_section(doc, resume, lang)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
